from docx import Document
import tempfile
import os


def generate_protocol_word(protocol_text: str, template_path: str = None) -> str:
    """
    Генерация документа Word по шаблону с сохранением форматирования.

    :param protocol_text: текст протокола, разбитый по строкам.
    :param template_path: путь к .docx-шаблону, из которого нужно взять
        форматирование. Если не указан, используется шаблон из папки templates.
    :return: путь к временному .docx-файлу.
    """
    # Если путь не передан — ищем шаблон в папке templates
    if template_path is None:
        base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        default_template = os.path.join(base_dir, 'templates', 'protocol_template.docx')
        if os.path.exists(default_template):
            template_path = default_template
        else:
            template_path = None

    # Загружаем шаблон или создаём пустой документ
    if template_path and os.path.exists(template_path):
        doc = Document(template_path)
        # Удалим все параграфы, если шаблон содержит заглушки
        for _ in range(len(doc.paragraphs)):
            p = doc.paragraphs[0]
            p.clear()
    else:
        doc = Document()

    # Добавляем текст с базовым стилем
    for line in protocol_text.split('\n'):
        if line.strip():  # пропускаем пустые строки
            doc.add_paragraph(line.strip(), style='Normal')
        else:
            doc.add_paragraph()

    # Сохраняем во временный файл
    with tempfile.NamedTemporaryFile('wb', delete=False, suffix='.docx') as tmp:
        doc.save(tmp.name)
        return tmp.name
